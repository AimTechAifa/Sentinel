"""Generate WORKFLOW.docx from workflow content."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "WORKFLOW.docx"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="E8EEF7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Release Desk — Workflow Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document explains how Sentinel / Release Desk fits together: what to do first, "
        "how pages connect, and the day-to-day steps for release managers."
    )

    doc.add_heading("What Release Desk does", level=1)
    doc.add_paragraph("Release Desk helps you:")
    add_numbered(doc, [
        "Plan what is releasing and when",
        "Prepare test environments for end-to-end validation",
        "Track upstream/downstream system dependencies",
        "Monitor portfolio health, P1 issues, and connector sync status",
    ])
    doc.add_paragraph(
        "All core data (departments, applications, environments, releases, bookings, mappings, versions) "
        "lives in the SQLite database. Pages read from the same source so changes in one place appear everywhere else."
    )

    doc.add_heading("Roles", level=1)
    add_table(doc, ["Role", "Access"], [
        ("Read only", "View dashboard, releases, calendar, bookings, mapping, versions"),
        ("Editor", "Everything read-only can do, plus create/edit releases, book environments, manage reference data, promote versions, sync connectors"),
        ("Admin", "Same as Editor (reference data and configuration)"),
    ])
    doc.add_paragraph("Sign in at /login and pick a role for the demo session.")

    doc.add_heading("High-level workflow", level=1)
    doc.add_paragraph(
        "1. Reference Data  →  2. Releases  →  3. Calendar\n"
        "                              ↓\n"
        "                    4. Env Booking  →  5. System Mapping  →  6. Versions & Config\n"
        "                              ↓                    ↓\n"
        "                         7. Dashboard  ←──────────────┘"
    )
    p = doc.add_paragraph()
    run = p.add_run("Rule of thumb: ")
    run.bold = True
    p.add_run("Set up master data once → plan releases → book and map environments → promote versions → use the dashboard for daily checks.")

    doc.add_heading("Step-by-step workflow", level=1)

    doc.add_heading("Step 0 — Start the app", level=2)
    doc.add_paragraph("npm run db:setup   # first time, or after schema changes")
    doc.add_paragraph("npm run dev")
    doc.add_paragraph("Open http://localhost:3000 and sign in as Editor or Admin for full interactivity.")

    doc.add_heading("Step 1 — Reference Data (foundation)", level=2)
    doc.add_paragraph("Page: Operations → Reference Data (/admin/reference-data)")
    doc.add_paragraph("Configure the master data every other page depends on:")
    add_table(doc, ["Section", "Fields"], [
        ("Departments", "name, head"),
        ("Applications", "name, department, type, product owner, tech lead, support, criticality"),
        ("Environments", "application, name, type (Dev/Test/UAT/Prod), owner, last DB refresh, status"),
    ])
    doc.add_paragraph("Actions:")
    add_bullets(doc, ["Add / edit / delete rows inline", "CSV upload for bulk import"])
    p = doc.add_paragraph()
    r = p.add_run("Why first? ")
    r.bold = True
    p.add_run("Dropdowns on Releases, Env Booking, and System Mapping all pull from these tables.")

    doc.add_heading("Step 2 — Releases (plan the work)", level=2)
    doc.add_paragraph("Page: Release Desk → Releases (/releases)")
    doc.add_paragraph("Each release includes:")
    add_bullets(doc, [
        "Release ID (releaseCode)",
        "Name",
        "Program / Project (use N/A for hotfixes or infra)",
        "Owner, status, release date",
        "Priority and impact (High / Medium / Low)",
        "Department",
        "Application(s)",
        "Depends on (other releases)",
    ])
    doc.add_paragraph("Actions:")
    add_bullets(doc, [
        "New release — create from the list page",
        "Click a row — open the release detail page",
        "Edit / Delete — from the list (Editor/Admin)",
    ])
    doc.add_paragraph("Detail page (/releases/[id]):")
    add_bullets(doc, [
        "Change status with quick buttons",
        "Record Go / No-Go (saved to audit trail)",
        "Add notes",
        "Jump to Dependencies, Env Booking, or System Mapping",
    ])

    doc.add_heading("Step 3 — Release Calendar (when things land)", level=2)
    doc.add_paragraph("Page: Release Desk → Calendar (/calendar)")
    add_bullets(doc, [
        "Period filter: Month | Quarter | Year",
        "View: Calendar grid or Timeline",
        "Releases come from the database (same as the Releases list)",
        "Click a release to open its detail page",
    ])
    doc.add_paragraph("Use this to see deployment windows and portfolio load for the selected period.")

    doc.add_heading("Step 4 — Environment Booking (can we test?)", level=2)
    doc.add_paragraph("Page: Release Desk → Env Booking (/booking)")
    doc.add_paragraph("Used when end-to-end testing needs one or more applications at the same time.")
    doc.add_paragraph("Workflow:")
    add_numbered(doc, [
        "Select applications from the multi-select dropdown (from reference data)",
        "Set From and To dates (calendar pickers)",
        "Enter a purpose (e.g. FIN SIT 1)",
        "Click Check Availability — if available, Book Now appears; if not, see who booked it",
        "Review Current bookings at the bottom of the page",
    ])
    doc.add_paragraph("Bookings are stored in the database and feed into System Mapping risk analysis.")

    doc.add_heading("Step 5 — System Mapping (how systems connect)", level=2)
    doc.add_paragraph("Page: Release Desk → System Mapping (/system-mapping)")
    doc.add_paragraph("Documents which application environment talks to which (upstream/downstream).")
    doc.add_paragraph("Workflow:")
    add_numbered(doc, [
        "Mapping notes — describe the setup in plain language",
        "Generate mapping from notes — AI suggests edges (or add mapping edge manually)",
        "Set analysis period (From / To dates)",
        "Review Mapping risks — flags when a required mapped environment is already booked",
        "Manage edges (delete custom edges as needed)",
    ])
    doc.add_paragraph(
        'Example risk: "SAP TEST is required by FIN UAT mapping but is booked by another team during your test window."'
    )

    doc.add_heading("Step 6 — Versions & Config (promotion and topology)", level=2)
    doc.add_paragraph("Page: Release Desk → Versions & Config (/environments)")
    add_bullets(doc, [
        "Release timeline — click entries to filter by department/app",
        "Environment booking cards — per application",
        "System topology — applications and environments",
        "Current Version matrix — DEV / TEST / PROD per application",
    ])
    doc.add_paragraph("Promotion workflow:")
    add_numbered(doc, [
        "Open the version matrix",
        "Find an application with drift (DEV/TEST ahead of PROD)",
        "Click Promote → to copy the version to the next stage (DEV→TEST or TEST→PROD)",
        "Requires Editor or Admin",
    ])

    doc.add_heading("Step 7 — Dashboard (daily monitoring)", level=2)
    doc.add_paragraph("Page: Dashboard (/dashboard)")
    add_numbered(doc, [
        "Period toggle — Month | Quarter | Year (default: Month)",
        "AI Daily Summary — portfolio briefing",
        "Summary counts — Planned, In progress, Blocked, At risk",
        "Connector last-sync — Jira, GitHub, ServiceNow, Confluence",
        "P1 issues only — items that may need hotfix attention",
    ])

    doc.add_heading("Step 8 — Connectors (integration health)", level=2)
    doc.add_paragraph("Page: Operations → Connectors (/connectors)")
    add_bullets(doc, [
        "Release Desk integrations panel — live sync times for the four MVP sources",
        "Sync now — refresh last-synced timestamp (Editor/Admin)",
        "Full connector catalog below is demo/static data for stakeholder presentations",
    ])

    doc.add_heading("End-to-end scenario", level=1)
    doc.add_paragraph("Goal: Ship FIN billing changes that depend on SAP and CRM for UAT.")
    add_table(doc, ["Step", "Where", "What you do"], [
        ("1", "Reference Data", "Confirm FIN, SAP, CRM apps and UAT/Test environments exist"),
        ("2", "Releases", "Create release RD-2026-0xxx, link FIN + SAP, set dependency on platform release"),
        ("3", "Calendar", "Confirm date fits the quarter plan"),
        ("4", "Env Booking", "Select FIN + SAP + CRM, pick UAT dates, Check → Book Now"),
        ("5", "System Mapping", "Set analysis period, verify FIN UAT → SAP TEST mapping, fix risks if booked"),
        ("6", "Versions & Config", "Promote SAP TEST build toward PROD when ready"),
        ("7", "Release detail", "Go/No-Go decision, status → In Progress / Complete"),
        ("8", "Dashboard", "Confirm no new P1s, review AI summary"),
    ])

    doc.add_heading("How data flows between pages", level=1)
    doc.add_paragraph(
        "Reference Data feeds Releases, Env Booking, System Mapping, and Versions & Config.\n"
        "Releases feed Calendar, Env Booking, Versions & Config, and Dashboard.\n"
        "Env Booking feeds System Mapping (risk analysis) and Dashboard.\n"
        "System Mapping feeds Dashboard.\n"
        "Connectors feed Dashboard (last-sync times)."
    )

    doc.add_heading("Page quick reference", level=1)
    add_table(doc, ["Route", "Purpose", "Interactive actions"], [
        ("/login", "Sign in", "Pick role (demo SSO)"),
        ("/dashboard", "Portfolio summary", "Period toggle, read AI summary & P1s"),
        ("/admin/reference-data", "Master data", "CRUD, CSV import"),
        ("/releases", "Release list", "Create, edit, delete, open detail"),
        ("/releases/[id]", "Release command", "Status, Go/No-Go, notes, audit trail"),
        ("/releases/[id]/dependencies", "Dependency graph", "Visual map of apps, deps, mapping"),
        ("/calendar", "Schedule view", "Period + Calendar/Timeline, open release"),
        ("/booking", "Env reservation", "Multi-app book, availability check"),
        ("/system-mapping", "Env relationships", "Notes, generate, add edges, risk scan"),
        ("/environments", "Versions & topology", "Promote versions, cross-filter panels"),
        ("/connectors", "Integrations", "Sync now (MVP sources)"),
    ])

    doc.add_heading("Demo vs production data", level=1)
    add_table(doc, ["Area", "Source"], [
        ("Reference data, releases, booking, mapping, versions, dashboard counts, P1s, connector sync (top 4)", "Database (MVP)"),
        ("Executive, Insights, Compare, Agents, Knowledge Graph, full connector catalog", "Demo / static data"),
        ("Authentication", "Demo role picker (real Microsoft SSO not wired)"),
    ])

    doc.add_heading("Suggested first-time path (15 minutes)", level=1)
    add_numbered(doc, [
        "Sign in as Editor",
        "Skim Reference Data (pre-seeded)",
        "Open Releases → click RD-2026-0140 → try status + Go",
        "Env Booking → select SAP + FIN → check dates → book",
        "System Mapping → set dates → review risks",
        "Versions & Config → promote SAP",
        "Dashboard → read AI summary and P1 list",
    ])
    doc.add_paragraph("That completes one full pass through the Release Desk workflow.")

    doc.add_paragraph()
    footer = doc.add_paragraph("Sentinel — Release Desk Workflow Guide")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    main()
